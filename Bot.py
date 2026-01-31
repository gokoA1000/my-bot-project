import logging
import os
import requests
import pdfkit
import cloudscraper
from telegram import Update, ReplyKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
from bs4 import BeautifulSoup
from docx import Document
from PIL import Image
from fpdf import FPDF
from markdownify import markdownify as md

# إعداد السجلات
logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s', level=logging.INFO)

TOKEN = os.getenv("BOT_TOKEN")

# لوحة المفاتيح الجديدة
keyboard = [['PDF', 'Word'], ['Text', 'Markdown']]
reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True)

user_data = {}

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "👋 مرحباً! البوت يعمل بنجاح.\n\n"
        "1️⃣ أرسل **رابط موقع** ثم اختر الصيغة (PDF, Word, Markdown..).\n"
        "2️⃣ أرسل **صورة** أو **ملف نصي** وسأحوله لـ PDF.",
        reply_markup=reply_markup
    )

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """معالجة الملفات والصور"""
    attachment = update.message.effective_attachment
    # التعامل مع الصور أو المستندات
    file_obj = await attachment[-1].get_file() if isinstance(attachment, list) else await attachment.get_file()
    
    file_path = "input_file"
    await file_obj.download_to_drive(file_path)

    try:
        # محاولة 1: هل هو صورة؟
        try:
            image = Image.open(file_path)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            image.save("output.pdf", "PDF", resolution=100.0)
            await update.message.reply_document(document=open("output.pdf", 'rb'), filename="image.pdf")
            os.remove("output.pdf")
            return
        except:
            pass # ليس صورة، ننتقل للتالي

        # محاولة 2: هل هو ملف نصي؟
        with open(file_path, "r", encoding="utf-8") as f:
            text_content = f.read()
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for line in text_content.split('\n'):
            safe_line = line.encode('latin-1', 'replace').decode('latin-1')
            pdf.cell(200, 10, txt=safe_line, ln=True)
        
        pdf.output("output.pdf")
        await update.message.reply_document(document=open("output.pdf", 'rb'), filename="text.pdf")
        os.remove("output.pdf")

    except Exception as e:
        await update.message.reply_text(f"❌ لم أستطع تحويل هذا الملف.\nالخطأ: {str(e)}")
    
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text
    chat_id = update.message.chat_id

    # استقبال الرابط
    if text.startswith('http'):
        user_data[chat_id] = text
        await update.message.reply_text(f"🔗 تم الحفظ: {text}\nاختر الصيغة 👇", reply_markup=reply_markup)
    
    # تنفيذ التحويل
    elif text in ['PDF', 'Word', 'Text', 'Markdown']:
        url = user_data.get(chat_id)
        if not url:
            await update.message.reply_text("⚠️ أرسل الرابط أولاً.")
            return

        await update.message.reply_text(f"⏳ جاري التحويل إلى {text}...")

        try:
            scraper = cloudscraper.create_scraper()
            response = scraper.get(url)
            
            if response.status_code != 200:
                await update.message.reply_text("❌ الموقع لا يستجيب.")
                return

            if text == 'PDF':
                path_wkhtmltopdf = '/usr/bin/wkhtmltopdf'
                config = pdfkit.configuration(wkhtmltopdf=path_wkhtmltopdf)
                options = {'enable-local-file-access': None}
                pdfkit.from_url(url, 'output.pdf', configuration=config, options=options)
                await update.message.reply_document(document=open('output.pdf', 'rb'), filename="web.pdf")
                os.remove('output.pdf')

            elif text == 'Word':
                soup = BeautifulSoup(response.text, 'lxml')
                doc = Document()
                doc.add_heading(url, 0)
                for p in soup.find_all('p'):
                    if p.text.strip():
                        doc.add_paragraph(p.text.strip())
                doc.save('output.docx')
                await update.message.reply_document(document=open('output.docx', 'rb'), filename="web.docx")
                os.remove('output.docx')

            elif text == 'Text':
                soup = BeautifulSoup(response.text, 'lxml')
                text_content = soup.get_text(separator='\n\n', strip=True)
                with open('output.txt', 'w', encoding='utf-8') as f:
                    f.write(text_content)
                await update.message.reply_document(document=open('output.txt', 'rb'), filename="web.txt")
                os.remove('output.txt')

            elif text == 'Markdown':
                md_text = md(response.text)
                with open('output.md', 'w', encoding='utf-8') as f:
                    f.write(md_text)
                await update.message.reply_document(document=open('output.md', 'rb'), filename="web.md")
                os.remove('output.md')

        except Exception as e:
            await update.message.reply_text(f"❌ حدث خطأ: {str(e)}")

def main():
    app = Application.builder().token(TOKEN).build()
    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    app.add_handler(MessageHandler(filters.PHOTO | filters.Document.ALL, handle_document))
    app.run_polling()

if __name__ == '__main__':
    main()
